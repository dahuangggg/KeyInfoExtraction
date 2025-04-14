import os
import sys
import argparse
import platform
from tqdm import tqdm
from docx import Document

def extract_text_from_docx(docx_path):
    """从docx文件中提取纯文本内容"""
    try:
        doc = Document(docx_path)
        full_text = []
        
        # 提取段落文本
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                full_text.append('\t'.join(row_text))
        
        return '\n'.join(full_text)
    except Exception as e:
        print(f"处理docx文件 {docx_path} 时出错: {e}")
        return None

def extract_text_from_doc(doc_path):
    """从doc文件中提取纯文本内容"""
    system = platform.system()
    
    # 方法1: 使用win32com (仅Windows系统)
    if system == "Windows":
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.visible = False
            doc = word.Documents.Open(os.path.abspath(doc_path))
            text = doc.Content.Text
            doc.Close()
            word.Quit()
            return text
        except ImportError:
            print("警告: win32com模块未安装，无法直接处理doc文件")
        except Exception as e:
            print(f"使用win32com处理doc文件 {doc_path} 时出错: {e}")
    
    # 方法2: 使用antiword (Linux/Mac系统)
    if system in ["Linux", "Darwin"]:
        try:
            import subprocess
            result = subprocess.run(['antiword', doc_path], capture_output=True, text=True)
            if result.returncode == 0:
                return result.stdout
            else:
                print(f"使用antiword处理文件失败: {result.stderr}")
        except FileNotFoundError:
            print("警告: antiword未安装，请使用命令安装: sudo apt-get install antiword (Ubuntu) 或 brew install antiword (Mac)")
        except Exception as e:
            print(f"使用antiword处理doc文件 {doc_path} 时出错: {e}")
    
    # 方法3: 尝试使用textract库
    try:
        import textract
        text = textract.process(doc_path).decode('utf-8')
        return text
    except ImportError:
        print("警告: textract模块未安装，请使用pip install textract安装")
    except Exception as e:
        print(f"使用textract处理doc文件 {doc_path} 时出错: {e}")
    
    print(f"无法处理doc文件 {doc_path}，请安装必要的依赖或将文件转换为docx格式")
    return None

def convert_doc_to_txt(doc_path, output_dir=None):
    """将doc/docx文件转换为txt文件"""
    # 确定文件类型
    is_docx = doc_path.lower().endswith('.docx')
    
    # 提取文本
    if is_docx:
        text_content = extract_text_from_docx(doc_path)
    else:
        text_content = extract_text_from_doc(doc_path)
    
    if text_content is None:
        return False
    
    # 确定输出路径
    base_name = os.path.basename(doc_path)
    file_name_without_ext = os.path.splitext(base_name)[0]
    
    if output_dir:
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        output_path = os.path.join(output_dir, f"{file_name_without_ext}.txt")
    else:
        output_path = os.path.join(os.path.dirname(doc_path), f"{file_name_without_ext}.txt")
    
    # 写入文本文件
    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text_content)
        print(f"已转换: {doc_path} -> {output_path}")
        return True
    except Exception as e:
        print(f"保存文件 {output_path} 时出错: {e}")
        return False

def batch_convert(input_dir, output_dir=None, recursive=False):
    """批量转换目录中的所有doc/docx文件"""
    success_count = 0
    fail_count = 0
    
    # 获取所有doc/docx文件
    doc_files = []
    if recursive:
        for root, _, files in os.walk(input_dir):
            for file in files:
                if file.lower().endswith(('.doc', '.docx')):
                    doc_files.append(os.path.join(root, file))
    else:
        doc_files = [os.path.join(input_dir, f) for f in os.listdir(input_dir) 
                    if f.lower().endswith(('.doc', '.docx')) and os.path.isfile(os.path.join(input_dir, f))]
    
    # 处理每个文件
    print(f"找到 {len(doc_files)} 个doc/docx文件需要转换")
    for doc_file in tqdm(doc_files, desc="转换进度"):
        if convert_doc_to_txt(doc_file, output_dir):
            success_count += 1
        else:
            fail_count += 1
    
    print(f"\n转换完成! 成功: {success_count}, 失败: {fail_count}")

def main():
    parser = argparse.ArgumentParser(description='将DOC/DOCX文件转换为TXT文件')
    parser.add_argument('input', help='输入的doc/docx文件或目录')
    parser.add_argument('-o', '--output', help='输出目录', default=None)
    parser.add_argument('-r', '--recursive', action='store_true', help='递归处理子目录')
    
    args = parser.parse_args()
    
    if os.path.isfile(args.input):
        if not args.input.lower().endswith(('.doc', '.docx')):
            print("输入文件必须是.doc或.docx格式")
            return
        convert_doc_to_txt(args.input, args.output)
    elif os.path.isdir(args.input):
        batch_convert(args.input, args.output, args.recursive)
    else:
        print(f"输入路径 {args.input} 不存在")

if __name__ == "__main__":
    main()